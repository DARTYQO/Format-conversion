import os
import sys
from pathlib import Path
from typing import Optional
from PyPDF2 import PdfReader, PdfWriter
import docx
from base import BaseConverter  # שינינו את הייבוא מיחסי למוחלט

class PDFConverter(BaseConverter):
    """Handles PDF file conversions."""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['docx', 'txt']
    
    def convert(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """
        Convert PDF to other formats.
        
        Args:
            input_path (Path): Path to input PDF file
            output_path (Path): Path to output file
            **kwargs: Additional conversion options
                - preserve_formatting (bool): Try to preserve document formatting
                
        Returns:
            bool: True if conversion successful, False otherwise
        """
        if not self.validate_input(input_path):
            return False
            
        try:
            preserve_formatting = kwargs.get('preserve_formatting', True)
            reader = PdfReader(str(input_path))
            
            if output_path.suffix.lower() == '.txt':
                return self._to_txt(reader, output_path)
            elif output_path.suffix.lower() == '.docx':
                return self._to_docx(reader, output_path, preserve_formatting)
            else:
                self.logger.error(f"Unsupported output format: {output_path.suffix}")
                return False
                
        except Exception as e:
            self.logger.error(f"PDF conversion failed: {str(e)}")
            return False
    
    def _to_txt(self, reader: PdfReader, output_path: Path) -> bool:
        """Convert PDF to TXT format."""
        try:
            with output_path.open('w', encoding='utf-8') as f:
                for page in reader.pages:
                    f.write(page.extract_text() + '\n')
            return True
        except Exception as e:
            self.logger.error(f"Failed to convert PDF to TXT: {str(e)}")
            return False
    
    def _to_docx(self, reader: PdfReader, output_path: Path, preserve_formatting: bool) -> bool:
        """Convert PDF to DOCX format."""
        try:
            doc = docx.Document()
            
            for page in reader.pages:
                text = page.extract_text()
                if preserve_formatting:
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                else:
                    doc.add_paragraph(text)
                
                # Add page break between pages
                if page != reader.pages[-1]:
                    doc.add_page_break()
            
            doc.save(str(output_path))
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to convert PDF to DOCX: {str(e)}")
            return False
