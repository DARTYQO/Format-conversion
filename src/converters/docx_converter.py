from pathlib import Path
from typing import Optional
import docx
from PyPDF2 import PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from .base import BaseConverter

class DOCXConverter(BaseConverter):
    """Handles DOCX file conversions."""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['pdf', 'txt']
    
    def convert(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """
        Convert DOCX to other formats.
        
        Args:
            input_path (Path): Path to input DOCX file
            output_path (Path): Path to output file
            **kwargs: Additional conversion options
                - preserve_formatting (bool): Try to preserve document formatting
                
        Returns:
            bool: True if conversion successful, False otherwise
        """
        if not self.validate_input(input_path):
            return False
            
        try:
            preserve_formatting = kwargs.get('preserve_formatting', True)
            doc = docx.Document(str(input_path))
            
            if output_path.suffix.lower() == '.txt':
                return self._to_txt(doc, output_path)
            elif output_path.suffix.lower() == '.pdf':
                return self._to_pdf(doc, output_path, preserve_formatting)
            else:
                self.logger.error(f"Unsupported output format: {output_path.suffix}")
                return False
                
        except Exception as e:
            self.logger.error(f"DOCX conversion failed: {str(e)}")
            return False
    
    def _to_txt(self, doc: docx.Document, output_path: Path) -> bool:
        """Convert DOCX to TXT format."""
        try:
            with output_path.open('w', encoding='utf-8') as f:
                for para in doc.paragraphs:
                    f.write(para.text + '\n')
                    
                # Handle tables
                for table in doc.tables:
                    for row in table.rows:
                        f.write('\t'.join(cell.text for cell in row.cells) + '\n')
            return True
        except Exception as e:
            self.logger.error(f"Failed to convert DOCX to TXT: {str(e)}")
            return False
    
    def _to_pdf(self, doc: docx.Document, output_path: Path, preserve_formatting: bool) -> bool:
        """Convert DOCX to PDF format."""
        try:
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            y = height - 72  # Start 1 inch from top
            
            for para in doc.paragraphs:
                if preserve_formatting:
                    # Try to preserve basic formatting
                    font_size = 12  # Default size
                    for run in para.runs:
                        if run.bold:
                            c.setFont("Helvetica-Bold", font_size)
                        elif run.italic:
                            c.setFont("Helvetica-Oblique", font_size)
                        else:
                            c.setFont("Helvetica", font_size)
                            
                        text = run.text
                        c.drawString(72, y, text)
                        y -= font_size + 2
                else:
                    # Simple text conversion
                    c.setFont("Helvetica", 12)
                    c.drawString(72, y, para.text)
                    y -= 14
                
                # New page if needed
                if y < 72:
                    c.showPage()
                    y = height - 72
            
            c.save()
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to convert DOCX to PDF: {str(e)}")
            return False
