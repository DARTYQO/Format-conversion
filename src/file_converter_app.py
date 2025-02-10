import os
import json
import threading
from PyQt5.QtWidgets import QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFileDialog, QComboBox, QToolBar, QStatusBar, QProgressBar, QMessageBox, QAction, QDialog, QFormLayout, QFrame, QCheckBox, QGroupBox, QListWidget, QListWidgetItem, QAbstractItemView, QMenu, QLineEdit, QTabWidget, QPushButton, QSpinBox
from PyQt5.QtCore import Qt, QTextCodec, QFileInfo
from PyQt5.QtGui import QIcon
from PyPDF2 import PdfReader, PdfMerger
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import docx
from drag_drop_list_widget import DragDropListWidget
from utils import STYLE_SHEET

class FileConverterApp(QMainWindow):
    def __init__(self):
        super().__init__()
        
        QTextCodec.setCodecForLocale(QTextCodec.codecForName('UTF-8'))
        
        self.output_folder = os.path.join(os.path.expanduser("~"), "Documents", "FileConverter")
        self.progress_bar = None
        self.format_combo = None
        self.output_dir = None
        self.last_output_dir = None
        
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)
        
        self.selected_files = []
        
        icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "icon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        self.init_ui()
        self.create_menus()
        self.create_toolbar()
        
        self.load_settings()
        
        self.setStyleSheet(STYLE_SHEET)

    def init_ui(self):
        self.setWindowTitle("Format conversion")
        self.setGeometry(100, 100, 1000, 600)
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(20)
        
        main_frame = QFrame()
        main_frame.setObjectName("mainFrame")
        main_frame.setFrameStyle(QFrame.StyledPanel)
        frame_layout = QHBoxLayout(main_frame)
        frame_layout.setContentsMargins(20, 20, 20, 20)
        frame_layout.setSpacing(20)
        
        self.tab_widget = QTabWidget()
        
        text_tab = QWidget()
        text_layout = QVBoxLayout(text_tab)
        
        files_group = QGroupBox("Selected Files")
        files_layout = QVBoxLayout()
        
        self.files_list = DragDropListWidget(self)
        self.files_list.setMinimumWidth(400)
        self.files_list.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.files_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.files_list.customContextMenuRequested.connect(self.show_context_menu)
        files_layout.addWidget(self.files_list)
        
        files_group.setLayout(files_layout)
        
        right_layout = QVBoxLayout()
        
        conversion_group = QGroupBox("Conversion Settings")
        conversion_layout = QVBoxLayout()
        
        format_layout = QFormLayout()
        self.format_combo = QComboBox()
        self.format_combo.addItems(['PDF', 'DOCX', 'TXT'])
        self.format_combo.setToolTip("Select target format")
        format_layout.addRow("Target Format", self.format_combo)
        
        self.preserve_format_cb = QCheckBox("Preserve Formatting")
        self.preserve_format_cb.setChecked(True)
        self.preserve_format_cb.setToolTip("Preserve the original document formatting")
        
        conversion_layout.addLayout(format_layout)
        conversion_layout.addWidget(self.preserve_format_cb)
        conversion_layout.addStretch()
        
        convert_btn = QPushButton("Convert")
        convert_btn.setMinimumHeight(50)
        convert_btn.setToolTip("Start file conversion")
        convert_btn.clicked.connect(self.convert_files)
        conversion_layout.addWidget(convert_btn)
        
        conversion_group.setLayout(conversion_layout)
        right_layout.addWidget(conversion_group)
        
        merge_split_group = QGroupBox("Merge/Split Settings")
        merge_split_layout = QVBoxLayout()
        
        action_layout = QFormLayout()
        self.action_combo = QComboBox()
        self.action_combo.addItems(['Merge All', 'Merge by Type', 'Split by Size', 'Split by Character Count'])
        self.action_combo.setToolTip("Select action type")
        self.action_combo.currentIndexChanged.connect(self.on_action_changed)
        action_layout.addRow("Action Type", self.action_combo)
        
        self.merge_type_combo = QComboBox()
        self.merge_type_combo.addItems(['PDF', 'DOCX', 'TXT'])
        self.merge_type_combo.setToolTip("Select file type for merging")
        self.merge_type_combo.setVisible(False)
        action_layout.addRow("File Type for Merging", self.merge_type_combo)
        
        self.split_value = QSpinBox()
        self.split_value.setRange(1, 999999)
        self.split_value.setValue(1000)
        self.split_value.setVisible(False)
        self.split_value_label = QLabel("Size in KB")
        self.split_value_label.setVisible(False)
        action_layout.addRow(self.split_value_label, self.split_value)
        
        merge_split_layout.addLayout(action_layout)
        merge_split_layout.addStretch()
        
        process_btn = QPushButton("Process")
        process_btn.setMinimumHeight(50)
        process_btn.setToolTip("Start merging/splitting files")
        process_btn.clicked.connect(self.process_files)
        merge_split_layout.addWidget(process_btn)
        
        merge_split_group.setLayout(merge_split_layout)
        right_layout.addWidget(merge_split_group)
        
        text_layout_main = QHBoxLayout()
        text_layout_main.addWidget(files_group)
        text_layout_main.addLayout(right_layout)
        text_layout.addLayout(text_layout_main)
        
        self.tab_widget.addTab(text_tab, "Text")
        
        audio_tab = QWidget()
        audio_layout = QVBoxLayout(audio_tab)
        audio_label = QLabel("Convert Audio Files")
        audio_drag_list = DragDropListWidget()
        audio_convert_button = QPushButton("Convert Audio Files")
        audio_layout.addWidget(audio_label)
        audio_layout.addWidget(audio_drag_list)
        audio_layout.addWidget(audio_convert_button)
        self.tab_widget.addTab(audio_tab, "Audio")
        
        video_tab = QWidget()
        video_layout = QVBoxLayout(video_tab)
        video_label = QLabel("Convert Video Files")
        video_drag_list = DragDropListWidget()
        video_convert_button = QPushButton("Convert Video Files")
        video_layout.addWidget(video_label)
        video_layout.addWidget(video_drag_list)
        video_layout.addWidget(video_convert_button)
        self.tab_widget.addTab(video_tab, "Video")
        
        code_tab = QWidget()
        code_layout = QVBoxLayout(code_tab)
        code_label = QLabel("Convert Code Files")
        code_drag_list = DragDropListWidget()
        code_convert_button = QPushButton("Convert Code Files")
        code_layout.addWidget(code_label)
        code_layout.addWidget(code_drag_list)
        code_layout.addWidget(code_convert_button)
        self.tab_widget.addTab(code_tab, "Code")
        
        frame_layout.addWidget(main_frame)
        main_layout.addWidget(self.tab_widget)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)
        
        self.status = QStatusBar()
        self.setStatusBar(self.status)
        
        self.status.showMessage("Ready")

    def create_menus(self):
        menubar = self.menuBar()
        
        file_menu = menubar.addMenu('File')
        
        open_action = QAction('Open File', self)
        open_action.setShortcut('Ctrl+O')
        open_action.triggered.connect(self.select_files)
        file_menu.addAction(open_action)
        
        add_folder_action = QAction("Add Folder", self)
        add_folder_action.triggered.connect(self.select_folder)
        file_menu.addAction(add_folder_action)
        
        open_output_dir_action = QAction('Open Output Directory', self)
        open_output_dir_action.setShortcut('Ctrl+D')
        open_output_dir_action.triggered.connect(self.open_output_directory)
        file_menu.addAction(open_output_dir_action)
        
        clear_list_action = QAction("Clear List", self)
        clear_list_action.triggered.connect(self.clear_file_list)
        file_menu.addAction(clear_list_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction('Exit', self)
        exit_action.setShortcut('Ctrl+Q')
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        theme_menu = menubar.addMenu("Theme")
        blue_theme_action = QAction("Blue", self)
        blue_theme_action.triggered.connect(lambda: self.change_theme("blue"))
        theme_menu.addAction(blue_theme_action)

        light_green_theme_action = QAction("Light Green", self)
        light_green_theme_action.triggered.connect(lambda: self.change_theme("light_green"))
        theme_menu.addAction(light_green_theme_action)

        yellow_theme_action = QAction("Yellow", self)
        yellow_theme_action.triggered.connect(lambda: self.change_theme("yellow"))
        theme_menu.addAction(yellow_theme_action)

        default_theme_action = QAction("Default", self)
        default_theme_action.triggered.connect(lambda: self.change_theme("default"))
        theme_menu.addAction(default_theme_action)
        
        settings_menu = menubar.addMenu("Settings")
        change_output_folder_action = QAction("Select Output Directory", self)
        change_output_folder_action.triggered.connect(self.show_settings_dialog)
        settings_menu.addAction(change_output_folder_action)
        
        help_menu = menubar.addMenu('Help')
        
        about_action = QAction('About', self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
        
        homepage_action = QAction('Homepage', self)
        homepage_action.triggered.connect(self.open_homepage)
        help_menu.addAction(homepage_action)

    def create_toolbar(self):
        pass

    def show_about(self):
        about_text = """
        File Converter - Version 2.0
        
        A software for converting, merging, and splitting files.
        Supports formats: PDF, DOCX, TXT
        
        Key Features:
        - Convert between different formats
        - Merge multiple files into one
        - Split files by size or character count
        - Drag and drop file support
        - Support for multiple text encodings
        - Innovative tab interface with support for:
          * Text file conversion
          * Audio file conversion
          * Video file conversion
          * Code file conversion
        """
        
        QMessageBox.about(self, "About File Converter", about_text)

    def open_homepage(self):
        import webbrowser
        webbrowser.open('https://github.com/DARTYQO/Format-conversion')

    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Folder")
        if folder:
            icon_provider = QFileIconProvider()
            for root, dirs, files in os.walk(folder):
                for file in files:
                    file_path = os.path.join(root, file)
                    if file_path.endswith(('.pdf', '.docx', '.txt')):
                        exists = False
                        for i in range(self.files_list.count()):
                            if self.files_list.item(i).data(Qt.UserRole) == file_path:
                                exists = True
                                break
                        
                        if not exists:
                            icon = icon_provider.icon(QFileInfo(file_path))
                            item = QListWidgetItem(icon, os.path.basename(file_path))
                            item.setData(Qt.UserRole, file_path)
                            self.files_list.addItem(item)
                            self.selected_files.append(file_path)
            
            self.status.showMessage(f"Files added from folder {folder}")

    def open_output_directory(self):
        output_dir = self.output_dir if hasattr(self, 'output_dir') and self.output_dir else self.output_folder
        if os.path.exists(output_dir):
            os.startfile(output_dir)
        else:
            QMessageBox.information(self, "Info", "No output directory selected or the directory does not exist")

    def open_file(self, file_path):
        try:
            os.startfile(file_path)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error opening file: {str(e)}")

    def show_context_menu(self, position):
        menu = QMenu()
        selected_items = self.files_list.selectedItems()
        
        if not selected_items:
            return
            
        open_action = menu.addAction("Open File")
        open_action.triggered.connect(lambda: self.open_file(selected_items[0].data(Qt.UserRole)))
        
        menu.addSeparator()
        
        convert_menu = menu.addMenu("Convert to...")
        pdf_action = convert_menu.addAction("PDF")
        pdf_action.triggered.connect(lambda: self.convert_selected_file(selected_items[0].data(Qt.UserRole), 'pdf'))
        
        docx_action = convert_menu.addAction("DOCX")
        docx_action.triggered.connect(lambda: self.convert_selected_file(selected_items[0].data(Qt.UserRole), 'docx'))
        
        txt_action = convert_menu.addAction("TXT")
        txt_action.triggered.connect(lambda: self.convert_selected_file(selected_items[0].data(Qt.UserRole), 'txt'))
        
        menu.addSeparator()
        
        if len(selected_items) > 1:
            merge_menu = menu.addMenu("Merge...")
            merge_all_action = merge_menu.addAction("Merge All")
            merge_all_action.triggered.connect(self.merge_all_files)
            
            merge_by_type_action = merge_menu.addAction("Merge by Type")
            merge_by_type_action.triggered.connect(self.merge_by_type)
        
        split_menu = menu.addMenu("Split...")
        split_size_action = split_menu.addAction("Split by Size")
        split_size_action.triggered.connect(lambda: self.split_by_size_single(selected_items[0].data(Qt.UserRole)))
        
        split_chars_action = split_menu.addAction("Split by Characters")
        split_chars_action.triggered.connect(lambda: self.split_by_chars_single(selected_items[0].data(Qt.UserRole)))
        
        menu.addSeparator()
        
        remove_action = menu.addAction("Remove")
        remove_action.triggered.connect(lambda: self.remove_selected_item(selected_items[0]))
        
        menu.exec_(self.files_list.mapToGlobal(position))

    def remove_selected_item(self, item):
        file_path = item.data(Qt.UserRole)
        self.files_list.takeItem(self.files_list.row(item))
        if file_path in self.selected_files:
            self.selected_files.remove(file_path)
        self.status.showMessage(f"The file {os.path.basename(file_path)} was removed from the list")

    def convert_selected_file(self, file_path, target_format):
        original_files = self.selected_files.copy()
        
        self.selected_files = [file_path]
        
        self.format_combo.setCurrentText(target_format.upper())
        
        self.convert_files()
        
        self.selected_files = original_files

    def split_by_size_single(self, file_path):
        original_files = self.selected_files.copy()
        
        self.selected_files = [file_path]
        
        self.split_by_size()
        
        self.selected_files = original_files

    def split_by_chars_single(self, file_path):
        original_files = self.selected_files.copy()
        
        self.selected_files = [file_path]
        
        self.split_by_chars()
        
        self.selected_files = original_files

    def convert_files(self):
        if not self.selected_files:
            QMessageBox.warning(self, "Error", "No files selected for conversion.")
            return
        
        self.progress_bar.setVisible(True)
        self.progress_bar.setMaximum(len(self.selected_files))
        self.progress_bar.setValue(0)
        
        target_format = self.format_combo.currentText().lower()
        
        conversion_thread = threading.Thread(target=self.run_conversion, args=(target_format,))
        conversion_thread.start()

    def run_conversion(self, target_format):
        QApplication.setOverrideCursor(Qt.WaitCursor)
        try:
            for i, file_path in enumerate(self.selected_files):
                try:
                    self.status.showMessage("Converting file...")
                    input_ext = os.path.splitext(file_path)[1].lower()
                    
                    if input_ext == '.pdf' and target_format in ['docx', 'txt']:
                        self.pdf_to_format(file_path, target_format)
                    elif input_ext == '.docx' and target_format in ['pdf', 'txt']:
                        self.docx_to_format(file_path, target_format)
                    elif input_ext == '.txt' and target_format in ['pdf', 'docx']:
                        self.txt_to_document(file_path, target_format)
                        
                    self.progress_bar.setValue(i + 1)
                    
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Error converting file: {str(e)}")
                    return
            
            self.progress_bar.setVisible(False)
            self.status.showMessage("Conversion completed successfully.")
        finally:
            QApplication.restoreOverrideCursor()

    def pdf_to_format(self, input_path, target_format):
        output_path = os.path.join(self.output_folder, 
                                  os.path.splitext(os.path.basename(input_path))[0] + f'.{target_format}')
        
        if target_format == 'docx':
            doc = docx.Document()
            reader = PdfReader(input_path)
            for page in reader.pages:
                doc.add_paragraph(page.extract_text())
            doc.save(output_path)
        elif target_format == 'txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                reader = PdfReader(input_path)
                for page in reader.pages:
                    f.write(page.extract_text() + '\n')

    def docx_to_format(self, input_path, target_format):
        output_path = os.path.join(self.output_folder, 
                                  os.path.splitext(os.path.basename(input_path))[0] + f'.{target_format}')
        
        if target_format == 'pdf':
            doc = docx.Document(input_path)
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            for para in doc.paragraphs:
                c.drawString(72, height - 72, para.text)
                height -= 12
                if height < 72:
                    c.showPage()
                    height = letter[1]
            c.save()
        elif target_format == 'txt':
            with open(input_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                for line in lines:
                    doc.add_paragraph(line.strip())
            doc.save(output_path)

    def load_settings(self):
        settings_path = os.path.join(self.output_folder, 'settings.json')
        if os.path.exists(settings_path):
            with open(settings_path, 'r', encoding='utf-8') as f:
                settings = json.load(f)
                self.output_folder = settings.get('output_folder', self.output_folder)

    def save_settings(self):
        settings_path = os.path.join(self.output_folder, 'settings.json')
        with open(settings_path, 'w', encoding='utf-8') as f:
            settings = {
                'output_folder': self.output_folder
            }
            json.dump(settings, f)

    def show_settings_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Settings")
        layout = QVBoxLayout()
        
        output_group = QGroupBox("Output Directory")
        output_layout = QHBoxLayout()
        
        self.output_dir_edit = QLineEdit()
        if hasattr(self, 'output_dir'):
            self.output_dir_edit.setText(self.output_dir)
        
        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(self.browse_output_dir)
        
        output_layout.addWidget(self.output_dir_edit)
        output_layout.addWidget(browse_btn)
        output_group.setLayout(output_layout)
        layout.addWidget(output_group)
        
        buttons = QHBoxLayout()
        ok_btn = QPushButton("OK")
        ok_btn.clicked.connect(lambda: self.save_settings(dialog))
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        
        buttons.addWidget(ok_btn)
        buttons.addWidget(cancel_btn)
        layout.addLayout(buttons)
        
        dialog.setLayout(layout)
        dialog.exec_()

    def browse_output_dir(self):
        dir_path = QFileDialog.getExistingDirectory(self, "Select Output Directory")
        if dir_path:
            self.output_dir_edit.setText(dir_path)

    def save_settings(self, dialog):
        self.output_dir = self.output_dir_edit.text()
        self.last_output_dir = self.output_dir
        dialog.accept()

    def clear_file_list(self):
        self.files_list.clear()
        self.selected_files.clear()
        self.status.showMessage("List cleared")

    def change_theme(self, theme_name):
        if theme_name == "blue":
            self.setStyleSheet("QMainWindow { background-color: #d0e7f9; }")
        elif theme_name == "light_green":
            self.setStyleSheet("QMainWindow { background-color: #d9f9d0; }")
        elif theme_name == "yellow":
            self.setStyleSheet("QMainWindow { background-color: #f9f9d0; }")
        else:
            self.setStyleSheet(STYLE_SHEET)

    def try_read_text_file(self, file_path):
        encodings = ['utf-8', 'utf-16', 'windows-1255', 'iso-8859-8', 'cp1255']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read(), encoding
            except UnicodeDecodeError:
                continue
            except Exception as e:
                continue
        
        raise ValueError(f"Cannot read the file {os.path.basename(file_path)} with any of the supported encodings")

    def merge_all_files(self):
        try:
            if not self.selected_files:
                QMessageBox.warning(self, "Error", "No files selected for merging")
                return

            file_types = set(os.path.splitext(f)[1].lower() for f in self.selected_files)
            filter_string = ""
            if '.pdf' in file_types:
                filter_string += "PDF Files (*.pdf);;"
            if '.docx' in file_types:
                filter_string += "Word Files (*.docx);;"
            if '.txt' in file_types:
                filter_string += "Text Files (*.txt);;"
            if not filter_string:
                QMessageBox.warning(self, "Error", "No suitable files for merging")
                return
            
            filter_string = filter_string.rstrip(';')

            output_file, selected_filter = QFileDialog.getSaveFileName(
                self,
                "Save Merged File",
                "",
                filter_string
            )
            
            if not output_file:
                return

            if output_file.endswith('.pdf'):
                merger = PdfMerger()
                pdf_files = [f for f in self.selected_files if f.endswith('.pdf')]
                if not pdf_files:
                    QMessageBox.warning(self, "Error", "No PDF files found for merging")
                    return
                
                for file in pdf_files:
                    try:
                        merger.append(file)
                    except Exception as e:
                        QMessageBox.critical(self, "Error", f"Error merging file {os.path.basename(file)}: {str(e)}")
                        return
                
                merger.write(output_file)
                merger.close()
                QMessageBox.information(self, "Success", "PDF files merged successfully")
                
            elif output_file.endswith('.docx'):
                docx_files = [f for f in self.selected_files if f.endswith('.docx')]
                if not docx_files:
                    QMessageBox.warning(self, "Error", "No DOCX files found for merging")
                    return
                
                self.merge_docx_files(docx_files, output_file)
                
            elif output_file.endswith('.txt'):
                txt_files = [f for f in self.selected_files if f.endswith('.txt')]
                if not txt_files:
                    QMessageBox.warning(self, "Error", "No text files found for merging")
                    return
                
                try:
                    with open(output_file, 'w', encoding='utf-8') as outfile:
                        for file in txt_files:
                            try:
                                content, _ = self.try_read_text_file(file)
                                outfile.write(content + '\n\n')
                            except Exception as e:
                                QMessageBox.critical(self, "Error", f"Error reading file {os.path.basename(file)}: {str(e)}")
                                return
                    QMessageBox.information(self, "Success", "Text files merged successfully")
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Error saving merged file: {str(e)}")
                    return
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Unexpected error: {str(e)}")

    def merge_docx_files(self, files, output_file):
        try:
            merged_doc = docx.Document()
            
            for i, file in enumerate(files):
                try:
                    current_doc = docx.Document(file)
                    
                    if i > 0:
                        merged_doc.add_page_break()
                    
                    for paragraph in current_doc.paragraphs:
                        new_paragraph = merged_doc.add_paragraph()
                        new_paragraph.alignment = paragraph.alignment
                        
                        for run in paragraph.runs:
                            new_run = new_paragraph.add_run(run.text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb
                    
                    for table in current_doc.tables:
                        new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = table.style
                        
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_cell = new_table.cell(i, j)
                                new_cell.text = cell.text
                                
                                if cell.paragraphs:
                                    for idx, paragraph in enumerate(cell.paragraphs):
                                        if idx == 0:
                                            new_paragraph = new_cell.paragraphs[0]
                                        else:
                                            new_paragraph = new_cell.add_paragraph()
                                        
                                        new_paragraph.alignment = paragraph.alignment
                                        for run in paragraph.runs:
                                            new_run = new_paragraph.add_run(run.text)
                                            new_run.bold = run.bold
                                            new_run.italic = run.italic
                                            new_run.underline = run.underline
                    
                    for shape in current_doc.inline_shapes:
                        if shape.type == 3:
                            try:
                                merged_doc.add_picture(shape._inline.graphic.graphicData.pic.blipFill.blip.embed)
                            except:
                                pass
                
                except Exception as e:
                    QMessageBox.warning(self, "Warning", f"Error merging file {os.path.basename(file)}: {str(e)}")
                    continue
            
            merged_doc.save(output_file)
            QMessageBox.information(self, "Success", "Documents merged successfully")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error merging documents: {str(e)}")
            return False
        
        return True

    def merge_by_type(self):
        try:
            selected_type = self.merge_type_combo.currentText().lower()
            files_to_merge = [file for file in self.selected_files if file.endswith(selected_type)]

            if not files_to_merge:
                QMessageBox.warning(self, "Error", f"No files of type {selected_type} found for merging")
                return

            output_file = QFileDialog.getSaveFileName(self, f"Save Merged {selected_type} File", "", f"{selected_type.upper()} Files (*.{selected_type})")[0]
            if not output_file:
                return

            if selected_type == 'pdf':
                merger = PdfMerger()
                for file in files_to_merge:
                    try:
                        merger.append(file)
                    except Exception as e:
                        QMessageBox.critical(self, "Error", f"Error merging file {os.path.basename(file)}: {str(e)}")
                        return
                
                try:
                    merger.write(output_file)
                    merger.close()
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Error saving merged file: {str(e)}")
                    return
                
            elif selected_type == 'docx':
                merged_doc = docx.Document()
                for file in files_to_merge:
                    try:
                        doc = docx.Document(file)
                        for paragraph in doc.paragraphs:
                            merged_doc.add_paragraph(paragraph.text)
                    except Exception as e:
                        QMessageBox.critical(self, "Error", f"Error merging file {os.path.basename(file)}: {str(e)}")
                        return
                
                try:
                    merged_doc.save(output_file)
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Error saving merged file: {str(e)}")
                    return
                
            elif selected_type == 'txt':
                try:
                    with open(output_file, 'w', encoding='utf-8') as outfile:
                        for file in files_to_merge:
                            try:
                                content, _ = self.try_read_text_file(file)
                                outfile.write(content + '\n')
                            except Exception as e:
                                QMessageBox.critical(self, "Error", f"Error reading file {os.path.basename(file)}: {str(e)}")
                                return
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Error saving merged file: {str(e)}")
                    return
            
            QMessageBox.information(self, "Success", "Files merged successfully")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Unexpected error: {str(e)}")

    def split_by_size(self):
        try:
            size_kb = self.split_value.value()
            for file in self.selected_files:
                if not file.endswith(('.pdf', '.docx', '.txt')):
                    continue

                base_name = os.path.splitext(file)[0]
                part_number = 1

                if file.endswith('.txt'):
                    try:
                        content, encoding = self.try_read_text_file(file)
                        current_size = 0
                        current_content = []
                        
                        for line in content.splitlines(True): 
                            line_size = len(line.encode(encoding))
                            if current_size + line_size > size_kb * 1024:
                                part_file = f"{base_name}_part{part_number}.txt"
                                try:
                                    with open(part_file, 'w', encoding=encoding) as outfile:
                                        outfile.writelines(current_content)
                                except Exception as e:
                                    QMessageBox.critical(self, "Error", f"Error saving part {part_number}: {str(e)}")
                                    return
                                
                                part_number += 1
                                current_size = line_size
                                current_content = [line]
                            else:
                                current_size += line_size
                                current_content.append(line)
                        
                        if current_content:
                            part_file = f"{base_name}_part{part_number}.txt"
                            try:
                                with open(part_file, 'w', encoding=encoding) as outfile:
                                    outfile.writelines(current_content)
                            except Exception as e:
                                QMessageBox.critical(self, "Error", f"Error saving last part: {str(e)}")
                                return
                            
                    except Exception as e:
                        QMessageBox.critical(self, "Error", f"Error reading file {os.path.basename(file)}: {str(e)}")
                        continue
            
            QMessageBox.information(self, "Success", "Files split successfully")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Unexpected error: {str(e)}")

    def split_by_chars(self):
        try:
            chars_count = self.split_value.value()
            for file in self.selected_files:
                if not file.endswith(('.pdf', '.docx', '.txt')):
                    continue

                base_name = os.path.splitext(file)[0]
                part_number = 1

                if file.endswith('.txt'):
                    try:
                        content, encoding = self.try_read_text_file(file)
                        current_chars = 0
                        current_content = []
                        
                        for line in content.splitlines(True): 
                            if current_chars + len(line) > chars_count:
                                part_file = f"{base_name}_part{part_number}.txt"
                                try:
                                    with open(part_file, 'w', encoding=encoding) as outfile:
                                        outfile.writelines(current_content)
                                except Exception as e:
                                    QMessageBox.critical(self, "Error", f"Error saving part {part_number}: {str(e)}")
                                    return
                                
                                part_number += 1
                                current_chars = len(line)
                                current_content = [line]
                            else:
                                current_chars += len(line)
                                current_content.append(line)
                        
                        if current_content:
                            part_file = f"{base_name}_part{part_number}.txt"
                            try:
                                with open(part_file, 'w', encoding=encoding) as outfile:
                                    outfile.writelines(current_content)
                            except Exception as e:
                                QMessageBox.critical(self, "Error", f"Error saving last part: {str(e)}")
                                return
                            
                    except Exception as e:
                        QMessageBox.critical(self, "Error", f"Error reading file {os.path.basename(file)}: {str(e)}")
                        continue
            
            QMessageBox.information(self, "Success", "Files split successfully")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Unexpected error: {str(e)}")

    def process_files(self):
        action = self.action_combo.currentText()
        if action == 'Merge All':
            self.merge_all_files()
        elif action == 'Merge by Type':
            self.merge_by_type()
        elif action == 'Split by Size':
            self.split_by_size()
        elif action == 'Split by Character Count':
            self.split_by_chars()

    def on_action_changed(self):
        action = self.action_combo.currentText()
        if action == 'Merge by Type':
            self.merge_type_combo.setVisible(True)
        else:
            self.merge_type_combo.setVisible(False)

        if action in ['Split by Size', 'Split by Character Count']:
            self.split_value.setVisible(True)
            self.split_value_label.setVisible(True)
        else:
            self.split_value.setVisible(False)
            self.split_value_label.setVisible(False)

    def select_files(self):
        try:
            files, _ = QFileDialog.getOpenFileNames(
                self,
                "Select Files",
                "",
                "Documents (*.pdf *.docx *.txt)"
            )
            
            if files:
                icon_provider = QFileIconProvider()
                for file_path in files:
                    exists = False
                    for i in range(self.files_list.count()):
                        if self.files_list.item(i).data(Qt.UserRole) == file_path:
                            exists = True
                            break
                    
                    if not exists:
                        icon = icon_provider.icon(QFileInfo(file_path))
                        item = QListWidgetItem(icon, os.path.basename(file_path))
                        item.setData(Qt.UserRole, file_path)
                        self.files_list.addItem(item)
                        self.selected_files.append(file_path)
                
                self.status.showMessage(f"{len(files)} files selected")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Unexpected error: {str(e)}")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = FileConverterApp()
    window.show()
    sys.exit(app.exec_())
