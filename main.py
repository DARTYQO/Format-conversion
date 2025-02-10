import sys
import os
import json
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                           QHBoxLayout, QPushButton, QLabel, QFileDialog,
                           QComboBox, QToolBar, QStatusBar, QProgressBar,
                           QStyle, QMessageBox, QAction, QDialog, QFormLayout,
                           QFrame, QCheckBox, QGroupBox, QListWidget, QListWidgetItem, QAbstractItemView, QMenu, QFileIconProvider, QSpinBox, QLineEdit, QTabWidget)
from PyQt5.QtCore import Qt, QSize, QFileInfo, QTextCodec
from PyQt5.QtGui import QIcon, QFont, QPalette, QColor
import docx
from docx import Document
from docx.shared import Inches
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import threading
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

STYLE_SHEET = """
QMainWindow {
    background-color: #f5f5f5;
}

QToolBar {
    background-color: transparent;
    border: none;
    padding: 0;
    spacing: 0;
}

QToolButton {
    background-color: transparent;
    border: none;
    padding: 0;
}

QStatusBar {
    background-color: #e0e0e0;
    color: #333333;
}

QPushButton {
    background-color: #ffffff;
    color: #333333;
    border: 1px solid #cccccc;
    border-radius: 4px;
    padding: 8px 16px;
    min-width: 100px;
}

QPushButton:hover {
    background-color: #e0e0e0;
}

QPushButton:pressed {
    background-color: #d0d0d0;
}

QComboBox {
    background-color: #ffffff;
    color: #333333;
    border: 1px solid #cccccc;
    border-radius: 4px;
    padding: 5px;
    min-width: 150px;
}

QListWidget {
    background-color: #f5f5f5;
    border: 1px solid #cccccc;
    border-radius: 4px;
    color: #333333;
}

QListWidget::item {
    border-bottom: 1px solid #cccccc;
}

QListWidget::item:selected {
    background-color: #e0e0e0;
}

QProgressBar {
    border-radius: 4px;
    background-color: #e0e0e0;
    text-align: center;
    color: #333333;
}

QProgressBar::chunk {
    background-color: #cccccc;
    border-radius: 4px;
}

QFrame#mainFrame {
    background-color: #f5f5f5;
    border: 1px solid #cccccc;
    border-radius: 8px;
}
"""

class DragDropListWidget(QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.parent = parent

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.CopyAction)
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event):
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.CopyAction)
            event.accept()
            icon_provider = QFileIconProvider()
            
            for url in event.mimeData().urls():
                file_path = url.toLocalFile()
                if file_path.endswith(('.pdf', '.docx', '.txt')):
                    # בדיקה אם הקובץ כבר קיים ברשימה
                    exists = False
                    for i in range(self.count()):
                        if self.item(i).data(Qt.UserRole) == file_path:
                            exists = True
                            break
                    
                    if not exists:
                        icon = icon_provider.icon(QFileInfo(file_path))
                        item = QListWidgetItem(icon, os.path.basename(file_path))
                        item.setData(Qt.UserRole, file_path)
                        self.addItem(item)
                        if self.parent:
                            self.parent.selected_files.append(file_path)
                            self.parent.status.showMessage(f"{self.count()} קבצים נבחרו")
        else:
            event.ignore()

class FileConverterApp(QMainWindow):
    def __init__(self):
        super().__init__()
        
        # הגדרת קידוד ברירת מחדל ל-UTF-8
        QTextCodec.setCodecForLocale(QTextCodec.codecForName('UTF-8'))
        
        # הגדרת תיקיית ברירת מחדל
        self.output_folder = os.path.join(os.path.expanduser("~"), "Documents", "FileConverter")
        self.progress_bar = None
        self.format_combo = None
        self.output_dir = None
        self.last_output_dir = None
        
        # יצירת תיקיית ברירת מחדל אם לא קיימת
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)
        
        # רשימת הקבצים שנבחרו
        self.selected_files = []
        
        # הגדרת אייקון לתוכנה
        icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "icon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        # קריאה לפונקציות הראשוניות
        self.init_ui()
        self.create_menus()
        self.create_toolbar()
        
        # הגדרת הממשק והטעינה
        # self.setup_i18n()
        self.load_settings()
        
        # הגדרת סגנון
        self.setStyleSheet(STYLE_SHEET)

    def init_ui(self):
        """הגדרת ממשק המשתמש הראשי"""
        self.setWindowTitle("Format conversion")
        self.setGeometry(100, 100, 1000, 600)
        
        # יצירת ווידג'ט מרכזי
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # יצירת layout ראשי
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(20)
        
        # יצירת מסגרת ראשית
        main_frame = QFrame()
        main_frame.setObjectName("mainFrame")
        main_frame.setFrameStyle(QFrame.StyledPanel)
        frame_layout = QHBoxLayout(main_frame)
        frame_layout.setContentsMargins(20, 20, 20, 20)
        frame_layout.setSpacing(20)
        
        # יצירת טאבים
        self.tab_widget = QTabWidget()
        
        # טאב טקסט (הקיים)
        text_tab = QWidget()
        text_layout = QVBoxLayout(text_tab)
        
        # אזור שמאל - רשימת קבצים
        files_group = QGroupBox("קבצים שנבחרו")
        files_layout = QVBoxLayout()
        
        self.files_list = DragDropListWidget(self)
        self.files_list.setMinimumWidth(400)
        self.files_list.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.files_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.files_list.customContextMenuRequested.connect(self.show_context_menu)
        files_layout.addWidget(self.files_list)
        
        files_group.setLayout(files_layout)
        
        # אזור ימין - הגדרות
        right_layout = QVBoxLayout()
        
        # הגדרות המרה
        conversion_group = QGroupBox("הגדרות המרה")
        conversion_layout = QVBoxLayout()
        
        # פורמט יעד
        format_layout = QFormLayout()
        self.format_combo = QComboBox()
        self.format_combo.addItems(['PDF', 'DOCX', 'TXT'])
        self.format_combo.setToolTip("בחר פורמט יעד")
        format_layout.addRow("פורמט יעד", self.format_combo)
        
        # הגדרות נוספות
        self.preserve_format_cb = QCheckBox("שמור על עיצוב")
        self.preserve_format_cb.setChecked(True)
        self.preserve_format_cb.setToolTip("שמור על עיצוב המסמך המקורי")
        
        conversion_layout.addLayout(format_layout)
        conversion_layout.addWidget(self.preserve_format_cb)
        conversion_layout.addStretch()
        
        # כפתור המרה
        convert_btn = QPushButton("המר")
        convert_btn.setMinimumHeight(50)
        convert_btn.setToolTip("התחל המרת קבצים")
        convert_btn.clicked.connect(self.convert_files)
        conversion_layout.addWidget(convert_btn)
        
        conversion_group.setLayout(conversion_layout)
        right_layout.addWidget(conversion_group)
        
        # הגדרות מיזוג ופיצול
        merge_split_group = QGroupBox("הגדרות מיזוג ופיצול")
        merge_split_layout = QVBoxLayout()
        
        # סוג פעולה
        action_layout = QFormLayout()
        self.action_combo = QComboBox()
        self.action_combo.addItems(['מיזוג הכל', 'מיזוג לפי סוג', 'פיצול לפי גודל', 'פיצול לפי מספר תווים'])
        self.action_combo.setToolTip("בחר סוג פעולה")
        self.action_combo.currentIndexChanged.connect(self.on_action_changed)
        action_layout.addRow("סוג פעולה", self.action_combo)
        
        # הגדרות נוספות למיזוג
        self.merge_type_combo = QComboBox()
        self.merge_type_combo.addItems(['PDF', 'DOCX', 'TXT'])
        self.merge_type_combo.setToolTip("בחר סוג קובץ למיזוג")
        self.merge_type_combo.setVisible(False)
        action_layout.addRow("סוג קובץ למיזוג", self.merge_type_combo)
        
        # הגדרות נוספות לפיצול
        self.split_value = QSpinBox()
        self.split_value.setRange(1, 999999)
        self.split_value.setValue(1000)
        self.split_value.setVisible(False)
        self.split_value_label = QLabel("גודל בKB")
        self.split_value_label.setVisible(False)
        action_layout.addRow(self.split_value_label, self.split_value)
        
        merge_split_layout.addLayout(action_layout)
        merge_split_layout.addStretch()
        
        # כפתור ביצוע
        process_btn = QPushButton("בצע")
        process_btn.setMinimumHeight(50)
        process_btn.setToolTip("התחל מיזוג/פיצול קבצים")
        process_btn.clicked.connect(self.process_files)
        merge_split_layout.addWidget(process_btn)
        
        merge_split_group.setLayout(merge_split_layout)
        right_layout.addWidget(merge_split_group)
        
        # סידור הלייאוט של טאב טקסט
        text_layout_main = QHBoxLayout()
        text_layout_main.addWidget(files_group)
        text_layout_main.addLayout(right_layout)
        text_layout.addLayout(text_layout_main)
        
        # הוספת הטאבים
        self.tab_widget.addTab(text_tab, "טקסט")
        
        # טאב אודיו
        audio_tab = QWidget()
        audio_layout = QVBoxLayout(audio_tab)
        audio_label = QLabel("המרת קבצי אודיו")
        audio_drag_list = DragDropListWidget()
        audio_convert_button = QPushButton("המר קבצי אודיו")
        audio_layout.addWidget(audio_label)
        audio_layout.addWidget(audio_drag_list)
        audio_layout.addWidget(audio_convert_button)
        self.tab_widget.addTab(audio_tab, "אודיו")
        
        # טאב וידאו
        video_tab = QWidget()
        video_layout = QVBoxLayout(video_tab)
        video_label = QLabel("המרת קבצי וידאו")
        video_drag_list = DragDropListWidget()
        video_convert_button = QPushButton("המר קבצי וידאו")
        video_layout.addWidget(video_label)
        video_layout.addWidget(video_drag_list)
        video_layout.addWidget(video_convert_button)
        self.tab_widget.addTab(video_tab, "וידאו")
        
        # טאב קוד
        code_tab = QWidget()
        code_layout = QVBoxLayout(code_tab)
        code_label = QLabel("המרת קבצי קוד")
        code_drag_list = DragDropListWidget()
        code_convert_button = QPushButton("המר קבצי קוד")
        code_layout.addWidget(code_label)
        code_layout.addWidget(code_drag_list)
        code_layout.addWidget(code_convert_button)
        self.tab_widget.addTab(code_tab, "קוד")
        
        # הוספת הטאבים לפריסה
        frame_layout.addWidget(main_frame)
        main_layout.addWidget(self.tab_widget)
        
        # סרגל התקדמות
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)
        
        # סרגל סטטוס
        self.status = QStatusBar()
        self.setStatusBar(self.status)
        
        self.status.showMessage("מוכן")

    def create_menus(self):
        """יצירת תפריטים"""
        menubar = self.menuBar()
        
        # תפריט קובץ
        file_menu = menubar.addMenu('קובץ')
        
        # פתיחת קובץ
        open_action = QAction('פתח קובץ', self)
        open_action.setShortcut('Ctrl+O')
        open_action.triggered.connect(self.select_files)
        file_menu.addAction(open_action)
        
        # פתיחת תיקייה
        add_folder_action = QAction("הוסף תיקייה", self)
        add_folder_action.triggered.connect(self.select_folder)
        file_menu.addAction(add_folder_action)
        
        # פתיחת תיקיית יעד
        open_output_dir_action = QAction('פתח תיקיית יעד', self)
        open_output_dir_action.setShortcut('Ctrl+D')
        open_output_dir_action.triggered.connect(self.open_output_directory)
        file_menu.addAction(open_output_dir_action)
        
        # ניקוי רשימה
        clear_list_action = QAction("נקה רשימה", self)
        clear_list_action.triggered.connect(self.clear_file_list)
        file_menu.addAction(clear_list_action)
        
        file_menu.addSeparator()
        
        # יציאה
        exit_action = QAction('יציאה', self)
        exit_action.setShortcut('Ctrl+Q')
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # תפריט ערכת נושא
        theme_menu = menubar.addMenu("ערכת נושא")
        blue_theme_action = QAction("תכלת", self)
        blue_theme_action.triggered.connect(lambda: self.change_theme("blue"))
        theme_menu.addAction(blue_theme_action)

        light_green_theme_action = QAction("ירוק בהיר", self)
        light_green_theme_action.triggered.connect(lambda: self.change_theme("light_green"))
        theme_menu.addAction(light_green_theme_action)

        yellow_theme_action = QAction("צהוב", self)
        yellow_theme_action.triggered.connect(lambda: self.change_theme("yellow"))
        theme_menu.addAction(yellow_theme_action)

        default_theme_action = QAction("רגיל", self)
        default_theme_action.triggered.connect(lambda: self.change_theme("default"))
        theme_menu.addAction(default_theme_action)
        
        # תפריט הגדרות
        settings_menu = menubar.addMenu("הגדרות")
        change_output_folder_action = QAction("בחירת תיקיית יעד", self)
        change_output_folder_action.triggered.connect(self.show_settings_dialog)
        settings_menu.addAction(change_output_folder_action)
        
        # תפריט עזרה
        help_menu = menubar.addMenu('עזרה')
        
        # אודות
        about_action = QAction('אודות', self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
        
        # אתר הבית
        homepage_action = QAction('אתר הבית', self)
        homepage_action.triggered.connect(self.open_homepage)
        help_menu.addAction(homepage_action)

    def create_toolbar(self):
        """יצירת סרגל כלים"""
        # הוסף את הפונקציה הזו מהקוד המקורי שלך
        pass

    def show_about(self):
        """הצגת חלון אודות"""
        about_text = """
        ממיר קבצים - גרסה 2.0
        
        תוכנה להמרה, מיזוג ופיצול של קבצים.
        תומכת בפורמטים: PDF, DOCX, TXT
        
        תכונות עיקריות:
        - המרה בין פורמטים שונים
        - מיזוג מספר קבצים לקובץ אחד
        - פיצול קבצים לפי גודל או מספר תווים
        - תמיכה בגרירה ושחרור של קבצים
        - תמיכה בקידודי טקסט שונים
        - ממשק טאבים חדשני עם תמיכה ב:
          * המרת קבצי טקסט
          * המרת קבצי אודיו
          * המרת קבצי וידאו
          * המרת קבצי קוד
        """
        
        QMessageBox.about(self, "אודות ממיר הקבצים", about_text)

    def open_homepage(self):
        """פתיחת אתר הבית של התוכנה"""
        import webbrowser
        webbrowser.open('https://github.com/DARTYQO/Format-conversion')  # החלף בכתובת האתר האמיתית
    def show_about(self):
        """הצגת חלון אודות"""
        about_text = """
        ממיר קבצים - גרסה 2.0
        
        תוכנה להמרה, מיזוג ופיצול של קבצים.
        תומכת בפורמטים: PDF, DOCX, TXT
        
        תכונות עיקריות:
        - המרה בין פורמטים שונים
        - מיזוג מספר קבצים לקובץ אחד
        - פיצול קבצים לפי גודל או מספר תווים
        - תמיכה בגרירה ושחרור של קבצים
        - תמיכה בקידודי טקסט שונים
        - ממשק טאבים חדשני עם תמיכה ב:
          * המרת קבצי טקסט
          * המרת קבצי אודיו
          * המרת קבצי וידאו
          * המרת קבצי קוד
        """
        
        QMessageBox.about(self, "אודות ממיר הקבצים", about_text)

    def open_homepage(self):
        """פתיחת אתר הבית של התוכנה"""
        import webbrowser
        webbrowser.open('https://github.com/DARTYQO/Format-conversion')  # החלף בכתובת האתר האמיתית

    def select_folder(self):
        """בחירת תיקייה והוספת כל הקבצים הנתמכים מתוכה"""
        folder = QFileDialog.getExistingDirectory(self, "בחר תיקייה")
        if folder:
            icon_provider = QFileIconProvider()
            for root, dirs, files in os.walk(folder):
                for file in files:
                    file_path = os.path.join(root, file)
                    if file_path.endswith(('.pdf', '.docx', '.txt')):
                        # בדיקה אם הקובץ כבר קיים ברשימה
                        exists = False
                        for i in range(self.files_list.count()):
                            if self.files_list.item(i).data(Qt.UserRole) == file_path:
                                exists = True
                                break
                        
                        if not exists:
                            icon = icon_provider.icon(QFileInfo(file_path))
                            item = QListWidgetItem(icon, os.path.basename(file_path))
                            item.setData(Qt.UserRole, file_path)
                            self.files_list.addItem(item)
                            self.selected_files.append(file_path)
            
            # עדכון סטטוס
            self.status.showMessage(f"נוספו קבצים מהתיקייה {folder}")

    def open_output_directory(self):
        """פתיחת תיקיית היעד"""
        output_dir = self.output_dir if hasattr(self, 'output_dir') and self.output_dir else self.output_folder
        if os.path.exists(output_dir):
            os.startfile(output_dir)
        else:
            QMessageBox.information(self, "מידע", "לא נבחרה תיקיית יעד או שהתיקייה אינה קיימת")

    def open_file(self, file_path):
        """פתיחת קובץ עם תוכנת ברירת המחדל"""
        try:
            os.startfile(file_path)
        except Exception as e:
            QMessageBox.critical(self, "שגיאה", f"שגיאה בפתיחת הקובץ: {str(e)}")

    def show_context_menu(self, position):
        """הצגת תפריט הקשר למקש ימני"""
        menu = QMenu()
        selected_items = self.files_list.selectedItems()
        
        if not selected_items:
            return
            
        # פתיחת קובץ
        open_action = menu.addAction("פתח קובץ")
        open_action.triggered.connect(lambda: self.open_file(selected_items[0].data(Qt.UserRole)))
        
        menu.addSeparator()
        
        # המרה לסוגי קבצים
        convert_menu = menu.addMenu("המר ל...")
        pdf_action = convert_menu.addAction("PDF")
        pdf_action.triggered.connect(lambda: self.convert_selected_file(selected_items[0].data(Qt.UserRole), 'pdf'))
        
        docx_action = convert_menu.addAction("DOCX")
        docx_action.triggered.connect(lambda: self.convert_selected_file(selected_items[0].data(Qt.UserRole), 'docx'))
        
        txt_action = convert_menu.addAction("TXT")
        txt_action.triggered.connect(lambda: self.convert_selected_file(selected_items[0].data(Qt.UserRole), 'txt'))
        
        menu.addSeparator()
        
        # מיזוג ופיצול
        if len(selected_items) > 1:
            merge_menu = menu.addMenu("מזג...")
            merge_all_action = merge_menu.addAction("מזג הכל")
            merge_all_action.triggered.connect(self.merge_all_files)
            
            merge_by_type_action = merge_menu.addAction("מזג לפי סוג")
            merge_by_type_action.triggered.connect(self.merge_by_type)
        
        split_menu = menu.addMenu("פצל...")
        split_size_action = split_menu.addAction("פצל לפי גודל")
        split_size_action.triggered.connect(lambda: self.split_by_size_single(selected_items[0].data(Qt.UserRole)))
        
        split_chars_action = split_menu.addAction("פצל לפי תווים")
        split_chars_action.triggered.connect(lambda: self.split_by_chars_single(selected_items[0].data(Qt.UserRole)))
        
        menu.addSeparator()
        
        # הסרת קובץ
        remove_action = menu.addAction("הסר")
        remove_action.triggered.connect(lambda: self.remove_selected_item(selected_items[0]))
        
        menu.exec_(self.files_list.mapToGlobal(position))

    def remove_selected_item(self, item):
        """הסרת פריט בודד מהרשימה"""
        file_path = item.data(Qt.UserRole)
        self.files_list.takeItem(self.files_list.row(item))
        if file_path in self.selected_files:
            self.selected_files.remove(file_path)
        self.status.showMessage(f"הקובץ {os.path.basename(file_path)} הוסר מהרשימה")

    def convert_selected_file(self, file_path, target_format):
        """המרת קובץ בודד לפורמט היעד"""
        # שמירת הקובץ המקורי
        original_files = self.selected_files.copy()
        
        # הגדרת הקובץ הנבחר כקובץ היחיד להמרה
        self.selected_files = [file_path]
        
        # הגדרת פורמט היעד
        self.format_combo.setCurrentText(target_format.upper())
        
        # ביצוע ההמרה
        self.convert_files()
        
        # שחזור רשימת הקבצים המקורית
        self.selected_files = original_files

    def split_by_size_single(self, file_path):
        """פיצול קובץ בודד לפי גודל"""
        # שמירת הקובץ המקורי
        original_files = self.selected_files.copy()
        
        # הגדרת הקובץ הנבחר כקובץ היחיד לפיצול
        self.selected_files = [file_path]
        
        # ביצוע הפיצול
        self.split_by_size()
        
        # שחזור רשימת הקבצים המקורית
        self.selected_files = original_files

    def split_by_chars_single(self, file_path):
        """פיצול קובץ בודד לפי מספר תווים"""
        # שמירת הקובץ המקורי
        original_files = self.selected_files.copy()
        
        # הגדרת הקובץ הנבחר כקובץ היחיד לפיצול
        self.selected_files = [file_path]
        
        # ביצוע הפיצול
        self.split_by_chars()
        
        # שחזור רשימת הקבצים המקורית
        self.selected_files = original_files

    def convert_files(self):
        if not self.selected_files:
            QMessageBox.warning(self, "שגיאה", "לא נבחרו קבצים להמרה.")
            return
        
        self.progress_bar.setVisible(True)
        self.progress_bar.setMaximum(len(self.selected_files))
        self.progress_bar.setValue(0)
        
        target_format = self.format_combo.currentText().lower()
        
        conversion_thread = threading.Thread(target=self.run_conversion, args=(target_format,))
        conversion_thread.start()

    def run_conversion(self, target_format):
        QApplication.setOverrideCursor(Qt.WaitCursor)
        try:
            for i, file_path in enumerate(self.selected_files):
                try:
                    self.status.showMessage("ממיר קובץ...")
                    input_ext = os.path.splitext(file_path)[1].lower()
                    
                    if input_ext == '.pdf' and target_format in ['docx', 'txt']:
                        self.pdf_to_format(file_path, target_format)
                    elif input_ext == '.docx' and target_format in ['pdf', 'txt']:
                        self.docx_to_format(file_path, target_format)
                    elif input_ext == '.txt' and target_format in ['pdf', 'docx']:
                        self.txt_to_document(file_path, target_format)
                        
                    self.progress_bar.setValue(i + 1)
                    
                except Exception as e:
                    QMessageBox.critical(self, "שגיאה", f"שגיאה בהמרת הקובץ: {str(e)}")
                    return
            
            self.progress_bar.setVisible(False)
            self.status.showMessage("ההמרה הושלמה בהצלחה.")
        finally:
            QApplication.restoreOverrideCursor()

    def pdf_to_format(self, input_path, target_format):
        output_path = os.path.join(self.output_folder, 
                                  os.path.splitext(os.path.basename(input_path))[0] + f'.{target_format}')
        
        if target_format == 'docx':
            # המרת PDF ל-DOCX
            doc = docx.Document()
            reader = PdfReader(input_path)
            for page in reader.pages:
                doc.add_paragraph(page.extract_text())
            doc.save(output_path)
        elif target_format == 'txt':
            # המרת PDF ל-TXT
            with open(output_path, 'w', encoding='utf-8') as f:
                reader = PdfReader(input_path)
                for page in reader.pages:
                    f.write(page.extract_text() + '\n')

    def docx_to_format(self, input_path, target_format):
        output_path = os.path.join(self.output_folder, 
                                  os.path.splitext(os.path.basename(input_path))[0] + f'.{target_format}')
        
        if target_format == 'pdf':
            # המרת DOCX ל-PDF
            doc = docx.Document(input_path)
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            for para in doc.paragraphs:
                c.drawString(72, height - 72, para.text)
                height -= 12
                if height < 72:
                    c.showPage()
                    height = letter[1]
            c.save()
        elif target_format == 'txt':
            # המרת DOCX ל-TXT
            with open(output_path, 'w', encoding='utf-8') as f:
                doc = docx.Document(input_path)
                for para in doc.paragraphs:
                    f.write(para.text + '\n')

    def txt_to_document(self, input_path, target_format):
        output_path = os.path.join(self.output_folder, 
                                  os.path.splitext(os.path.basename(input_path))[0] + f'.{target_format}')
        
        if target_format == 'pdf':
            # המרת TXT ל-PDF
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            with open(input_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                for line in lines:
                    c.drawString(72, height - 72, line.strip())
                    height -= 12
                    if height < 72:
                        c.showPage()
                        height = letter[1]
            c.save()
        elif target_format == 'docx':
            # המרת TXT ל-DOCX
            doc = docx.Document()
            with open(input_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                for line in lines:
                    doc.add_paragraph(line.strip())
            doc.save(output_path)

    def load_settings(self):
        settings_path = os.path.join(self.output_folder, 'settings.json')
        if os.path.exists(settings_path):
            with open(settings_path, 'r', encoding='utf-8') as f:
                settings = json.load(f)
                self.output_folder = settings.get('output_folder', self.output_folder)

    def save_settings(self):
        settings_path = os.path.join(self.output_folder, 'settings.json')
        with open(settings_path, 'w', encoding='utf-8') as f:
            settings = {
                'output_folder': self.output_folder
            }
            json.dump(settings, f)

    def show_settings_dialog(self):
        """הצגת חלון הגדרות"""
        dialog = QDialog(self)
        dialog.setWindowTitle("הגדרות")
        layout = QVBoxLayout()
        
        # בחירת תיקיית יעד
        output_group = QGroupBox("תיקיית יעד")
        output_layout = QHBoxLayout()
        
        self.output_dir_edit = QLineEdit()
        if hasattr(self, 'output_dir'):
            self.output_dir_edit.setText(self.output_dir)
        
        browse_btn = QPushButton("עיון...")
        browse_btn.clicked.connect(self.browse_output_dir)
        
        output_layout.addWidget(self.output_dir_edit)
        output_layout.addWidget(browse_btn)
        output_group.setLayout(output_layout)
        layout.addWidget(output_group)
        
        # כפתורי אישור וביטול
        buttons = QHBoxLayout()
        ok_btn = QPushButton("אישור")
        ok_btn.clicked.connect(lambda: self.save_settings(dialog))
        cancel_btn = QPushButton("ביטול")
        cancel_btn.clicked.connect(dialog.reject)
        
        buttons.addWidget(ok_btn)
        buttons.addWidget(cancel_btn)
        layout.addLayout(buttons)
        
        dialog.setLayout(layout)
        dialog.exec_()

    def browse_output_dir(self):
        """בחירת תיקיית יעד"""
        dir_path = QFileDialog.getExistingDirectory(self, "בחר תיקיית יעד")
        if dir_path:
            self.output_dir_edit.setText(dir_path)

    def save_settings(self, dialog):
        """שמירת הגדרות"""
        self.output_dir = self.output_dir_edit.text()
        self.last_output_dir = self.output_dir  # שמירת תיקיית היעד האחרונה
        dialog.accept()

    def clear_file_list(self):
        """ניקוי רשימת הקבצים"""
        self.files_list.clear()
        self.selected_files.clear()
        self.status.showMessage("הרשימה נוקתה")

    def change_theme(self, theme_name):
        if theme_name == "blue":
            self.setStyleSheet("QMainWindow { background-color: #d0e7f9; }")
        elif theme_name == "light_green":
            self.setStyleSheet("QMainWindow { background-color: #d9f9d0; }")
        elif theme_name == "yellow":
            self.setStyleSheet("QMainWindow { background-color: #f9f9d0; }")
        else:
            self.setStyleSheet(STYLE_SHEET)

    def try_read_text_file(self, file_path):
        """נסיון לקרוא קובץ טקסט עם קידודים שונים"""
        encodings = ['utf-8', 'utf-16', 'windows-1255', 'iso-8859-8', 'cp1255']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read(), encoding
            except UnicodeDecodeError:
                continue
            except Exception as e:
                continue
        
        raise ValueError(f"לא ניתן לקרוא את הקובץ {os.path.basename(file_path)} עם אף אחד מהקידודים הנתמכים")

    def merge_all_files(self):
        try:
            # בדיקה שיש קבצים ברשימה
            if not self.selected_files:
                QMessageBox.warning(self, "שגיאה", "לא נבחרו קבצים למיזוג")
                return

            # קביעת סוג הקובץ לפי הקבצים שנבחרו
            file_types = set(os.path.splitext(f)[1].lower() for f in self.selected_files)
            filter_string = ""
            if '.pdf' in file_types:
                filter_string += "PDF Files (*.pdf);;"
            if '.docx' in file_types:
                filter_string += "Word Files (*.docx);;"
            if '.txt' in file_types:
                filter_string += "Text Files (*.txt);;"
            if not filter_string:
                QMessageBox.warning(self, "שגיאה", "אין קבצים מתאימים למיזוג")
                return
            
            # הסרת ;; מהסוף אם קיים
            filter_string = filter_string.rstrip(';')

            # פתיחת דיאלוג שמירת קובץ עם הפילטרים המתאימים
            output_file, selected_filter = QFileDialog.getSaveFileName(
                self,
                "שמור קובץ ממוזג",
                "",
                filter_string
            )
            
            if not output_file:
                return

            if output_file.endswith('.pdf'):
                merger = PdfMerger()
                pdf_files = [f for f in self.selected_files if f.endswith('.pdf')]
                if not pdf_files:
                    QMessageBox.warning(self, "שגיאה", "לא נמצאו קבצי PDF למיזוג")
                    return
                
                for file in pdf_files:
                    try:
                        merger.append(file)
                    except Exception as e:
                        QMessageBox.critical(self, "שגיאה", f"שגיאה במיזוג הקובץ {os.path.basename(file)}: {str(e)}")
                        return
                
                merger.write(output_file)
                merger.close()
                QMessageBox.information(self, "הצלחה", "קבצי ה-PDF מוזגו בהצלחה")
                
            elif output_file.endswith('.docx'):
                docx_files = [f for f in self.selected_files if f.endswith('.docx')]
                if not docx_files:
                    QMessageBox.warning(self, "שגיאה", "לא נמצאו קבצי DOCX למיזוג")
                    return
                
                self.merge_docx_files(docx_files, output_file)
                
            elif output_file.endswith('.txt'):
                txt_files = [f for f in self.selected_files if f.endswith('.txt')]
                if not txt_files:
                    QMessageBox.warning(self, "שגיאה", "לא נמצאו קבצי טקסט למיזוג")
                    return
                
                try:
                    with open(output_file, 'w', encoding='utf-8') as outfile:
                        for file in txt_files:
                            try:
                                content, _ = self.try_read_text_file(file)
                                outfile.write(content + '\n\n')  # הוספת שתי שורות ריקות בין הקבצים
                            except Exception as e:
                                QMessageBox.critical(self, "שגיאה", f"שגיאה בקריאת הקובץ {os.path.basename(file)}: {str(e)}")
                                return
                    QMessageBox.information(self, "הצלחה", "קבצי הטקסט מוזגו בהצלחה")
                except Exception as e:
                    QMessageBox.critical(self, "שגיאה", f"שגיאה בשמירת הקובץ הממוזג: {str(e)}")
                    return
            
        except Exception as e:
            QMessageBox.critical(self, "שגיאה", f"שגיאה לא צפויה: {str(e)}")

    def merge_docx_files(self, files, output_file):
        """מיזוג קבצי Word"""
        try:
            # יצירת מסמך חדש
            merged_doc = Document()
            
            for i, file in enumerate(files):
                try:
                    # פתיחת המסמך הנוכחי
                    current_doc = Document(file)
                    
                    # אם זה לא הקובץ הראשון, הוסף מעבר עמוד
                    if i > 0:
                        merged_doc.add_page_break()
                    
                    # העתקת כל הפסקאות עם העיצוב שלהן
                    for paragraph in current_doc.paragraphs:
                        # העתקת הפסקה עם העיצוב
                        new_paragraph = merged_doc.add_paragraph()
                        new_paragraph.alignment = paragraph.alignment
                        
                        # העתקת הטקסט עם העיצוב
                        for run in paragraph.runs:
                            new_run = new_paragraph.add_run(run.text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb
                    
                    # העתקת טבלאות
                    for table in current_doc.tables:
                        new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = table.style
                        
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_cell = new_table.cell(i, j)
                                new_cell.text = cell.text
                                
                                # העתקת עיצוב התא
                                if cell.paragraphs:
                                    for idx, paragraph in enumerate(cell.paragraphs):
                                        if idx == 0:
                                            new_paragraph = new_cell.paragraphs[0]
                                        else:
                                            new_paragraph = new_cell.add_paragraph()
                                        
                                        new_paragraph.alignment = paragraph.alignment
                                        for run in paragraph.runs:
                                            new_run = new_paragraph.add_run(run.text)
                                            new_run.bold = run.bold
                                            new_run.italic = run.italic
                                            new_run.underline = run.underline
                    
                    # העתקת תמונות
                    for shape in current_doc.inline_shapes:
                        if shape.type == 3:  # תמונה
                            try:
                                merged_doc.add_picture(shape._inline.graphic.graphicData.pic.blipFill.blip.embed)
                            except:
                                pass  # אם יש בעיה בהעתקת תמונה, נמשיך הלאה
                
                except Exception as e:
                    QMessageBox.warning(self, "אזהרה", f"שגיאה במיזוג הקובץ {os.path.basename(file)}: {str(e)}")
                    continue
            
            # שמירת המסמך הממוזג
            merged_doc.save(output_file)
            QMessageBox.information(self, "הצלחה", "המסמכים מוזגו בהצלחה")
            
        except Exception as e:
            QMessageBox.critical(self, "שגיאה", f"שגיאה במיזוג המסמכים: {str(e)}")
            return False
        
        return True

    def merge_by_type(self):
        try:
            selected_type = self.merge_type_combo.currentText().lower()
            files_to_merge = [file for file in self.selected_files if file.endswith(selected_type)]

            if not files_to_merge:
                QMessageBox.warning(self, "שגיאה", f"אין קבצים מסוג {selected_type} למיזוג")
                return

            output_file = QFileDialog.getSaveFileName(self, f"שמור קובץ {selected_type} ממוזג", "", f"{selected_type.upper()} Files (*.{selected_type})")[0]
            if not output_file:
                return

            if selected_type == 'pdf':
                merger = PdfMerger()
                for file in files_to_merge:
                    try:
                        merger.append(file)
                    except Exception as e:
                        QMessageBox.critical(self, "שגיאה", f"שגיאה במיזוג הקובץ {os.path.basename(file)}: {str(e)}")
                        return
                
                try:
                    merger.write(output_file)
                    merger.close()
                except Exception as e:
                    QMessageBox.critical(self, "שגיאה", f"שגיאה בשמירת הקובץ הממוזג: {str(e)}")
                    return
                
            elif selected_type == 'docx':
                merged_doc = Document()
                for file in files_to_merge:
                    try:
                        doc = Document(file)
                        for paragraph in doc.paragraphs:
                            merged_doc.add_paragraph(paragraph.text)
                    except Exception as e:
                        QMessageBox.critical(self, "שגיאה", f"שגיאה במיזוג הקובץ {os.path.basename(file)}: {str(e)}")
                        return
                
                try:
                    merged_doc.save(output_file)
                except Exception as e:
                    QMessageBox.critical(self, "שגיאה", f"שגיאה בשמירת הקובץ הממוזג: {str(e)}")
                    return
                
            elif selected_type == 'txt':
                try:
                    with open(output_file, 'w', encoding='utf-8') as outfile:
                        for file in files_to_merge:
                            try:
                                content, _ = self.try_read_text_file(file)
                                outfile.write(content + '\n')
                            except Exception as e:
                                QMessageBox.critical(self, "שגיאה", f"שגיאה בקריאת הקובץ {os.path.basename(file)}: {str(e)}")
                                return
                except Exception as e:
                    QMessageBox.critical(self, "שגיאה", f"שגיאה בשמירת הקובץ הממוזג: {str(e)}")
                    return
            
            QMessageBox.information(self, "הצלחה", "המיזוג הושלם בהצלחה")
            
        except Exception as e:
            QMessageBox.critical(self, "שגיאה", f"שגיאה לא צפויה: {str(e)}")

    def split_by_size(self):
        try:
            size_kb = self.split_value.value()
            for file in self.selected_files:
                if not file.endswith(('.pdf', '.docx', '.txt')):
                    continue

                base_name = os.path.splitext(file)[0]
                part_number = 1

                if file.endswith('.txt'):
                    try:
                        content, encoding = self.try_read_text_file(file)
                        current_size = 0
                        current_content = []
                        
                        for line in content.splitlines(True):  # keepends=True to keep newlines
                            line_size = len(line.encode(encoding))
                            if current_size + line_size > size_kb * 1024:
                                # Write current chunk
                                part_file = f"{base_name}_part{part_number}.txt"
                                try:
                                    with open(part_file, 'w', encoding=encoding) as outfile:
                                        outfile.writelines(current_content)
                                except Exception as e:
                                    QMessageBox.critical(self, "שגיאה", f"שגיאה בשמירת חלק {part_number}: {str(e)}")
                                    return
                                
                                part_number += 1
                                current_size = line_size
                                current_content = [line]
                            else:
                                current_size += line_size
                                current_content.append(line)
                        
                        # Write last chunk if exists
                        if current_content:
                            part_file = f"{base_name}_part{part_number}.txt"
                            try:
                                with open(part_file, 'w', encoding=encoding) as outfile:
                                    outfile.writelines(current_content)
                            except Exception as e:
                                QMessageBox.critical(self, "שגיאה", f"שגיאה בשמירת החלק האחרון: {str(e)}")
                                return
                            
                    except Exception as e:
                        QMessageBox.critical(self, "שגיאה", f"שגיאה בקריאת הקובץ {os.path.basename(file)}: {str(e)}")
                        continue
                # TODO: Add splitting logic for PDF and DOCX if needed
            
            QMessageBox.information(self, "הצלחה", "הפיצול הושלם בהצלחה")
            
        except Exception as e:
            QMessageBox.critical(self, "שגיאה", f"שגיאה לא צפויה: {str(e)}")

    def split_by_chars(self):
        try:
            chars_count = self.split_value.value()
            for file in self.selected_files:
                if not file.endswith(('.pdf', '.docx', '.txt')):
                    continue

                base_name = os.path.splitext(file)[0]
                part_number = 1

                if file.endswith('.txt'):
                    try:
                        content, encoding = self.try_read_text_file(file)
                        current_chars = 0
                        current_content = []
                        
                        for line in content.splitlines(True):  # keepends=True to keep newlines
                            if current_chars + len(line) > chars_count:
                                # Write current chunk
                                part_file = f"{base_name}_part{part_number}.txt"
                                try:
                                    with open(part_file, 'w', encoding=encoding) as outfile:
                                        outfile.writelines(current_content)
                                except Exception as e:
                                    QMessageBox.critical(self, "שגיאה", f"שגיאה בשמירת חלק {part_number}: {str(e)}")
                                    return
                                
                                part_number += 1
                                current_chars = len(line)
                                current_content = [line]
                            else:
                                current_chars += len(line)
                                current_content.append(line)
                        
                        # Write last chunk if exists
                        if current_content:
                            part_file = f"{base_name}_part{part_number}.txt"
                            try:
                                with open(part_file, 'w', encoding=encoding) as outfile:
                                    outfile.writelines(current_content)
                            except Exception as e:
                                QMessageBox.critical(self, "שגיאה", f"שגיאה בשמירת החלק האחרון: {str(e)}")
                                return
                            
                    except Exception as e:
                        QMessageBox.critical(self, "שגיאה", f"שגיאה בקריאת הקובץ {os.path.basename(file)}: {str(e)}")
                        continue
                # TODO: Add splitting logic for PDF and DOCX if needed
            
            QMessageBox.information(self, "הצלחה", "הפיצול הושלם בהצלחה")
            
        except Exception as e:
            QMessageBox.critical(self, "שגיאה", f"שגיאה לא צפויה: {str(e)}")

    def process_files(self):
        action = self.action_combo.currentText()
        if action == 'מיזוג הכל':
            self.merge_all_files()
        elif action == 'מיזוג לפי סוג':
            self.merge_by_type()
        elif action == 'פיצול לפי גודל':
            self.split_by_size()
        elif action == 'פיצול לפי מספר תווים':
            self.split_by_chars()

    def on_action_changed(self):
        action = self.action_combo.currentText()
        if action == 'מיזוג לפי סוג':
            self.merge_type_combo.setVisible(True)
        else:
            self.merge_type_combo.setVisible(False)

        if action in ['פיצול לפי גודל', 'פיצול לפי מספר תווים']:
            self.split_value.setVisible(True)
            self.split_value_label.setVisible(True)
        else:
            self.split_value.setVisible(False)
            self.split_value_label.setVisible(False)

    def select_files(self):
        """בחירת קבצים להמרה"""
        try:
            files, _ = QFileDialog.getOpenFileNames(
                self,
                "בחר קבצים",
                "",
                "Documents (*.pdf *.docx *.txt)"
            )
            
            if files:
                icon_provider = QFileIconProvider()
                for file_path in files:
                    # בדיקה אם הקובץ כבר קיים ברשימה
                    exists = False
                    for i in range(self.files_list.count()):
                        if self.files_list.item(i).data(Qt.UserRole) == file_path:
                            exists = True
                            break
                    
                    if not exists:
                        icon = icon_provider.icon(QFileInfo(file_path))
                        item = QListWidgetItem(icon, os.path.basename(file_path))
                        item.setData(Qt.UserRole, file_path)
                        self.files_list.addItem(item)
                        self.selected_files.append(file_path)
                
                # עדכון סטטוס
                self.status.showMessage(f"{len(files)} קבצים נבחרו")
        except Exception as e:
            QMessageBox.critical(self, "שגיאה", f"שגיאה בלתי צפויה: {str(e)}")

    def show_about(self):
        """הצגת חלון אודות"""
        about_text = """
        ממיר קבצים - גרסה 1.0
        
        תוכנה להמרה, מיזוג ופיצול של קבצים.
        תומכת בפורמטים: PDF, DOCX, TXT
        
        תכונות עיקריות:
        - המרה בין פורמטים שונים
        - מיזוג מספר קבצים לקובץ אחד
        - פיצול קבצים לפי גודל או מספר תווים
        - תמיכה בגרירה ושחרור של קבצים
        - תמיכה בקידודי טקסט שונים
        """
        
        QMessageBox.about(self, "אודות ממיר הקבצים", about_text)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = FileConverterApp()
    window.show()
    sys.exit(app.exec_())
